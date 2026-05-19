"""Document extraction service.

Pure functions over raw bytes; no I/O beyond the input. Each extractor returns
a (summary, chunks) tuple wrapped in an ExtractedDocument.

Supported types:
  pdf   — PyMuPDF (fitz)
  docx  — python-docx
  xlsx  — openpyxl
  png   — Pillow (extract dimensions + basic OCR-free description; for the
          hackathon we surface the image's dimensions and let the model see
          alt text. Real OCR would plug in Tesseract.)
  image — generic Pillow path for jpg/jpeg
"""
from __future__ import annotations

import io
from dataclasses import dataclass


@dataclass
class ExtractedDocument:
    summary: str
    chunks: list[str]


class DocumentExtractionError(Exception):
    """Raised when a document cannot be extracted."""


class DocumentService:
    """Sync extraction service — called from FastAPI worker thread."""

    def extract(
        self, *, kind: str, content: bytes, filename: str = ""
    ) -> ExtractedDocument:
        kind = kind.lower()
        if kind == "pdf":
            return self._extract_pdf(content, filename)
        if kind == "docx":
            return self._extract_docx(content, filename)
        if kind == "xlsx":
            return self._extract_xlsx(content, filename)
        if kind in ("image", "png"):
            return self._extract_image(content, filename)
        raise DocumentExtractionError(f"unsupported kind: {kind}")

    # -------------------------------------------- pdf

    @staticmethod
    def _extract_pdf(content: bytes, filename: str) -> ExtractedDocument:
        try:
            import fitz  # type: ignore[import-not-found]
        except ImportError as e:  # pragma: no cover
            raise DocumentExtractionError("PyMuPDF not installed") from e

        try:
            with fitz.open(stream=content, filetype="pdf") as pdf:
                chunks: list[str] = []
                for page_no, page in enumerate(pdf, start=1):
                    text = (page.get_text() or "").strip()
                    if not text:
                        continue
                    # Chunk by page first; further split if very long.
                    chunks.append(f"[page {page_no}] {text[:4000]}")
                if not chunks:
                    return ExtractedDocument(
                        summary=f"{filename or 'document.pdf'}: empty or image-only PDF",
                        chunks=[],
                    )
                summary = _summarise(chunks, filename or "document.pdf", "PDF")
                return ExtractedDocument(summary=summary, chunks=chunks)
        except Exception as e:
            raise DocumentExtractionError(f"pdf extraction failed: {e}") from e

    # -------------------------------------------- docx

    @staticmethod
    def _extract_docx(content: bytes, filename: str) -> ExtractedDocument:
        try:
            from docx import Document as DocxDoc  # type: ignore[import-not-found]
        except ImportError as e:  # pragma: no cover
            raise DocumentExtractionError("python-docx not installed") from e
        try:
            doc = DocxDoc(io.BytesIO(content))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            # Tables
            for tbl in doc.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            chunks = paragraphs
            if not chunks:
                return ExtractedDocument(
                    summary=f"{filename or 'document.docx'}: empty document",
                    chunks=[],
                )
            summary = _summarise(chunks, filename or "document.docx", "DOCX")
            return ExtractedDocument(summary=summary, chunks=chunks)
        except Exception as e:
            raise DocumentExtractionError(f"docx extraction failed: {e}") from e

    # -------------------------------------------- xlsx

    @staticmethod
    def _extract_xlsx(content: bytes, filename: str) -> ExtractedDocument:
        try:
            from openpyxl import load_workbook  # type: ignore[import-not-found]
        except ImportError as e:  # pragma: no cover
            raise DocumentExtractionError("openpyxl not installed") from e
        try:
            wb = load_workbook(io.BytesIO(content), data_only=True, read_only=True)
            chunks: list[str] = []
            for ws in wb.worksheets:
                rows_extracted: list[str] = []
                row_count = 0
                for row in ws.iter_rows(values_only=True):
                    row_count += 1
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        rows_extracted.append(" | ".join(cells)[:300])
                    if row_count >= 200:
                        break
                if rows_extracted:
                    chunks.append(f"[sheet {ws.title}] " + " ;; ".join(rows_extracted[:100]))
            wb.close()
            if not chunks:
                return ExtractedDocument(
                    summary=f"{filename or 'document.xlsx'}: empty workbook",
                    chunks=[],
                )
            summary = _summarise(chunks, filename or "document.xlsx", "XLSX")
            return ExtractedDocument(summary=summary, chunks=chunks)
        except Exception as e:
            raise DocumentExtractionError(f"xlsx extraction failed: {e}") from e

    # -------------------------------------------- image

    @staticmethod
    def _extract_image(content: bytes, filename: str) -> ExtractedDocument:
        try:
            from PIL import Image  # type: ignore[import-not-found]
        except ImportError as e:  # pragma: no cover
            raise DocumentExtractionError("Pillow not installed") from e
        try:
            with Image.open(io.BytesIO(content)) as img:
                w, h = img.size
                mode = img.mode
                fmt = img.format or "?"
            summary = (
                f"{filename or 'image'}: {fmt} image, {w}x{h}, mode={mode}. "
                f"Visual content not transcribed in this build."
            )
            return ExtractedDocument(summary=summary, chunks=[summary])
        except Exception as e:
            raise DocumentExtractionError(f"image extraction failed: {e}") from e


def _summarise(chunks: list[str], filename: str, kind_label: str) -> str:
    """Build a short text summary from extracted chunks for prompt context."""
    head = chunks[0]
    if len(head) > 500:
        head = head[:500] + "…"
    chars = sum(len(c) for c in chunks)
    return (
        f"{filename} ({kind_label}, {len(chunks)} chunk(s), {chars} chars). "
        f"Excerpt: {head}"
    )
