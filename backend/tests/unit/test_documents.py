"""Unit tests for DocumentService extractors."""
from __future__ import annotations

import io

import pytest

from app.documents.service import (
    DocumentExtractionError,
    DocumentService,
    ExtractedDocument,
)


def _png_bytes(w: int = 16, h: int = 16) -> bytes:
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (w, h), color="red").save(buf, format="PNG")
    return buf.getvalue()


def _jpg_bytes(w: int = 16, h: int = 16) -> bytes:
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (w, h), color="blue").save(buf, format="JPEG")
    return buf.getvalue()


def _docx_bytes(paragraphs: list[str]) -> bytes:
    from docx import Document

    d = Document()
    for p in paragraphs:
        d.add_paragraph(p)
    # Also add a small table
    tbl = d.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "A1"
    tbl.cell(0, 1).text = "B1"
    tbl.cell(1, 0).text = "A2"
    tbl.cell(1, 1).text = "B2"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _xlsx_bytes() -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Header A", "Header B", "Header C"])
    ws.append([1, 2, 3])
    ws.append([4, 5, 6])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _pdf_bytes(text: str = "Hello PDF World") -> bytes:
    """Build a minimal PDF using reportlab so PyMuPDF can extract its text."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.drawString(72, 800, text)
    c.showPage()
    c.save()
    return buf.getvalue()


# ---------------------------------------------------------------- pdf


def test_extract_pdf_returns_text():
    svc = DocumentService()
    res = svc.extract(kind="pdf", content=_pdf_bytes("ATRIO board meeting"), filename="x.pdf")
    assert isinstance(res, ExtractedDocument)
    assert any("ATRIO" in c for c in res.chunks)
    assert "ATRIO" in res.summary or "PDF" in res.summary


def test_extract_pdf_empty_returns_empty_chunks():
    svc = DocumentService()
    # An empty PDF (no content) — generate then strip text
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    # No content
    c.showPage()
    c.save()
    res = svc.extract(kind="pdf", content=buf.getvalue(), filename="empty.pdf")
    # Empty PDF produces empty chunks
    assert res.chunks == [] or all("page" in c for c in res.chunks)


def test_extract_pdf_invalid_bytes_raises():
    svc = DocumentService()
    with pytest.raises(DocumentExtractionError, match="pdf extraction failed"):
        svc.extract(kind="pdf", content=b"not a pdf at all", filename="x.pdf")


# ---------------------------------------------------------------- docx


def test_extract_docx_returns_paragraphs_and_table():
    svc = DocumentService()
    body = _docx_bytes(["Para 1", "Para 2", "Final paragraph"])
    res = svc.extract(kind="docx", content=body, filename="d.docx")
    assert "Para 1" in res.chunks
    assert any("A1 | B1" in c for c in res.chunks)
    assert "DOCX" in res.summary


def test_extract_docx_empty():
    """A docx with no paragraphs and no tables produces empty chunks."""
    from docx import Document as DocxDoc

    d = DocxDoc()
    buf = io.BytesIO()
    d.save(buf)
    svc = DocumentService()
    res = svc.extract(kind="docx", content=buf.getvalue(), filename="d.docx")
    assert res.chunks == []
    assert "empty" in res.summary


def test_extract_docx_invalid_raises():
    svc = DocumentService()
    with pytest.raises(DocumentExtractionError, match="docx"):
        svc.extract(kind="docx", content=b"not a docx", filename="d.docx")


# ---------------------------------------------------------------- xlsx


def test_extract_xlsx_returns_rows():
    svc = DocumentService()
    res = svc.extract(kind="xlsx", content=_xlsx_bytes(), filename="s.xlsx")
    assert any("Header A" in c for c in res.chunks)
    assert "XLSX" in res.summary


def test_extract_xlsx_invalid_raises():
    svc = DocumentService()
    with pytest.raises(DocumentExtractionError, match="xlsx"):
        svc.extract(kind="xlsx", content=b"not xlsx", filename="x.xlsx")


# ---------------------------------------------------------------- image


def test_extract_png_returns_dimensions():
    svc = DocumentService()
    res = svc.extract(kind="png", content=_png_bytes(20, 30), filename="i.png")
    assert "PNG" in res.summary
    assert "20x30" in res.summary


def test_extract_jpg_returns_dimensions():
    svc = DocumentService()
    res = svc.extract(kind="image", content=_jpg_bytes(40, 50), filename="i.jpg")
    assert "JPEG" in res.summary
    assert "40x50" in res.summary


def test_extract_image_invalid_raises():
    svc = DocumentService()
    with pytest.raises(DocumentExtractionError, match="image"):
        svc.extract(kind="image", content=b"not an image at all" * 4, filename="i.jpg")


# ---------------------------------------------------------------- unsupported


def test_extract_unsupported_kind_raises():
    svc = DocumentService()
    with pytest.raises(DocumentExtractionError, match="unsupported kind"):
        svc.extract(kind="weird-format", content=b"x", filename="x.xx")
